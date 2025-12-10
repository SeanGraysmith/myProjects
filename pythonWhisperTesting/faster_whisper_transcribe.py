#This is my faster-whisper transcription implementation, notably will ignore silence better than regular whisper
#Creates and adds timestamps at each segment of speech in form HR:MN:SC
#Only implements turbo model, work extremely well for webinar style audio samples. 
#Relies upon installation of cuda, faster-whisper, probably only works on windows, may need to troubleshot other instllations before it works.
#https://pypi.org/project/faster-whisper/0.3.0/ for installation and reference
#
#AUTHOR: Sean Graysmith
print(" ██████╗   ██████╗ \n"+"██╔════╝  ██╔════╝ \n"+"╚█████╗   ██║  ███╗\n"+" ╚═══██╗  ██║   ██║\n"+"██████╔╝  ╚██████╔╝\n"+"╚═════╝    ╚═════╝ ")


from faster_whisper import WhisperModel
from docx import Document
import os

def format_timestamp(seconds: float):
    hrs = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = seconds % 60
    if(hrs>0):
        return f"{hrs:02d}:{mins:02d}:{secs:05.2f}"
    else:
        return f"{mins:02d}:{secs:05.2f}"

#initializing stuff and things
audioName = ""
heading = ""
modelName = "turbo" #hard-coded for now
currentDirectory = os.getcwd()

print("Loading model " + modelName,end="...",flush=True)
model = WhisperModel(modelName,device="cuda",compute_type="float16")
print("done")

print("Creating and initializing document",end="...",flush=True)
document = Document()
print("done")

#user inpurt
audioName = input("Provide the name of the audio file to transcribe including the extension: ")
heading = input("Provide the title/heading to the .docx to be created: ")

#customize this heading as needed
print("Adding heading to document",end="...",flush=True)
document.add_heading(heading, level=0)
print("done")


#oh babyyyy
print("Transcribing audio file",end="...",flush=True)
segments, info = model.transcribe(audioName,
                          language="en",
                          task="transcribe",
                          beam_size=5,
                          best_of=5,
                          vad_filter=True,
                          vad_parameters={"min_silence_duration_ms":1000},
                          temperature=0.3,
                          condition_on_previous_text=False)
print("done")

print("Finalizing document and saving to: " + currentDirectory,end="...",flush=True)
full_transcription=""
for seg in segments:
    start = format_timestamp(seg.start)
    end = format_timestamp(seg.end)
    full_transcription += f"[{start} → {end}] {seg.text}\n"
document.add_paragraph(full_transcription)
document.save(audioName+"_transcribed.docx")
print("done..Thanks!")