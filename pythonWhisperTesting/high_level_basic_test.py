#This is a simple script that automates the use of whisper to transcribe a audio file
#The outfile is the name of the audio file appended with _transcribed.docx
#Currently only runs in the current directory..more to come soon?
#
#
#AUTHOR: Sean Graysmith

import whisper
from docx import Document
import os

#initializing stuff and things
audioName = ""
heading = ""
modelName = "turbo" #hard-coded for now
currentDirectory = os.getcwd()

print("Loading model " + modelName,end="...",flush=True)
model = whisper.load_model(modelName)
print("done")

print("Creating and initializing document",end="...",flush=True)
document = Document()
print("done")

#user inpurt
audioName = input("Provide the name of the audio file to transcribe including the extension: ")
heading = input("Provide the title/heading to the .docx to be created: ")

#customize this heading as needed
print("Adding heading to document",end="...",flush=True)
document.add_heading(heading, level=0)
print("done")


#oh babyyyy
print("Transcribing audio file",end="...",flush=True)
result = model.transcribe(audioName)
print("done")

print("Finalizing document and saving to: " + currentDirectory,end="...",flush=True)
document.add_paragraph(result["text"])
document.save(filename+"_transcribed.docx")
print("done..Thanks!")